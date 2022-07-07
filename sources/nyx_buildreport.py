
import os
import re
import sys
import docx
import json
import datetime
import cachetools
import matplotlib
import collections
import numpy as np
import pandas as pd
matplotlib.use('TkAgg')
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from dateutil.parser import parse
from elastic_helper import es_helper
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from elasticsearch import Elasticsearch as ES, RequestsHttpConnection as RC

#######################################################################################
# move_table_after
#######################################################################################
def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)

#######################################################################################
# add_hyperlink
#######################################################################################
def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

#######################################################################################
# insert_paragraph_after
#######################################################################################
def insert_paragraph_after(inparagraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    inparagraph._p.addnext(new_p)
    new_para = Paragraph(new_p, inparagraph._parent)
    if text is not None:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para

#######################################################################################
# create_table
#######################################################################################
def create_table( paragraph,df, style=None,title=""):
    global template
    table = template.add_table(df.shape[0]+1, df.shape[1])

    if style:
        table.style = style
        
    table_cells = table._cells

    for i in range(0,df.shape[0]+1):
        row_cells = table_cells[i*df.shape[1]:(i+1)*df.shape[1]]

        if i==0:

            for idx, col in enumerate(df.columns):
                row_cells[idx].text = col
        else:
            for idx, col in enumerate(df.columns):
                row_cells[idx].text = str(df.loc[(int(i-1)), col])


    table.autofit = True
    paragraph.text=title
    move_table_after(table, paragraph)

    return table

#######################################################################################
# replaceText
#######################################################################################
def replaceText(replacementHT,text):
    for keypair in replacementHT:
        try:
            text=text.replace("${"+keypair+"}",str(replacementHT[keypair])).replace('.0 %',' %')
        except:
            pass
    return text

#######################################################################################
# Start
#######################################################################################

print("Starting v0.1")
localmode=False
try:
    os.environ["LOCAL_MODE"]
    localmode=True
except:
    pass

try:
    report=json.loads(sys.argv[1])
except:
    report={"id": "id_19485666_27963523", "creds": {"token": "933921e1-58ae-4948-a941-037e9ce2e915", "user": {"filters": [], "firstname": "Arnaud", "id": "amarchand@icloud.com", "language": "en", "lastname": "Marchand", "login": "amarchand", "password": "", "phone": "0033497441962", "privileges": ["admin"], "user": "amarchand@icloud.com"}}, "report": {"description": "My First Report", "exec": "report1", "generatePDF": True, "icon": "file", "output": ["txt", "pdf"], "parameters": [{"name": "param1", "title": "Param1", "type": "text", "value": "test"}, {"name": "param2", "title": "Param2", "type": "interval", "value": ["2019-11-20T23:00:00.000Z", "2019-11-21T23:00:00.000Z"]}], "privileges": [], "reportType": "python", "title": "My First Report"}, "privileges": ["admin"], "treatment": {"status": "Waiting", "creation": "2019-11-22T08:24:03.288Z", "start": "2019-11-22T09:24:03.293651"}, "@timestamp": "2019-11-22T08:24:03.288Z", "output": "/opt/sources/generated/id_19485666_27963523"}

print("==== Args")
print(report)
print("==== Args")
print(report["report"]["parameters"])

params={}

for param in report["report"]["parameters"]:
    if param["type"]=="interval":
        params[param["name"]+"_start"]=parse(param["value"][0])
        params[param["name"]+"_end"]=parse(param["value"][1])
    else:
        params[param["name"]]=param["value"]

if not localmode:
    prepath = "./reports/notebooks/"
else:
    prepath = "./reports/notebooks/"

# ELASTIC SEARCH

if not localmode:
    host_params="http://esnodebal:9200"
    es = ES(hosts=[host_params])
    print(es.info())
else:        
    print("Connect to elastic via password")
    host_params = {'host':os.environ["ELK_URL"], 'port':int(os.environ["ELK_PORT"]), 'use_ssl':True}
    
    es = ES([host_params], connection_class=RC, http_auth=(os.environ["ELK_LOGIN"], os.environ["ELK_PASSWORD"]),  use_ssl=True ,verify_certs=False)
    print(es.info())

docpath=prepath+report["report"]["notebook"]+".docx"
ipynbpath=prepath+report["report"]["notebook"]+".ipynb"

# DOCS

if os.path.isfile(docpath) :
    print("DOC (%s) found." %(docpath))
else:
    print("ERROR DOC (%s) found." %(docpath))
    sys.exit(1) 

template = Document(docpath)

# PYTHON
if os.path.isfile(ipynbpath) :
    print("IPYNB (%s) found." %(docpath))
else:
    print("ERROR IPYNB (%s) found." %(docpath))
    sys.exit(1) 



replacementHT={}
replacementHT.update(params)

# READ notebook

reportfunctions={}

with open(ipynbpath, 'r') as content_file:
    content = content_file.read()
    jsoncontent=json.loads(content)
    for cell in jsoncontent["cells"]:
        if cell["cell_type"]=="code":
            if len(cell["source"])>0 and "#@ONLOAD" in cell["source"][0]:
                newcode="".join(cell["source"])
                exec(newcode)
            if len(cell["source"])>0 and "#@PARAGRAPH=" in cell["source"][0]:
                newcode="".join(cell["source"])
                reportfunctions["${"+cell["source"][0].replace("#@PARAGRAPH=","").strip()+"}"]=newcode
            

# FILL table cells
pattern = r"\${.*}"

for table in template.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:                            
                if re.search(pattern, paragraph.text):  
                    match= re.search(pattern, paragraph.text).group(0)
                    print(">>>>> FOUND CELL :%s" %(match))  
                    if match in reportfunctions:
                        paragraph.text=paragraph.text.replace(match,"")
                        exec(reportfunctions[match])
                    else:
                        paragraph.text=replaceText(replacementHT,paragraph.text)                   

# Fill paragraphs

for paragraph in template.paragraphs:     
    if re.search(pattern, paragraph.text):                
        match= re.search(pattern, paragraph.text).group(0)
        print(">>>>> FOUND PARAGRAPH :%s" %(match))                    
        paragraph.text=replaceText(replacementHT,paragraph.text)        
        if match in reportfunctions:
            paragraph.text=paragraph.text.replace(match,"")
            exec(reportfunctions[match])


# Fill headers
if not "noHeader" in replacementHT:
    for section in template.sections:
        header = section.header
        if re.search(pattern, header.text):                
            match= re.search(pattern, header.text).group(0)
            print(">>>>> FOUND PARAGRAPH :%s" %(match))                    
            header.text=replaceText(replacementHT,header.text)        
            if match in reportfunctions:
                header.text=header.text.replace(match,"")
                exec(reportfunctions[match])

# Fill footer
if not "noFooter" in replacementHT:
    for section in template.sections:
        footer = section.footer
        if re.search(pattern, footer.text):                
            match= re.search(pattern, footer.text).group(0)
            print(">>>>> FOUND PARAGRAPH :%s" %(match))                    
            footer.text=replaceText(replacementHT,footer.text)        
            if match in reportfunctions:
                footer.text=footer.text.replace(match,"")
                exec(reportfunctions[match])            



# SAVE REPORT

if not localmode:
    template.save(report["output"]+".docx")
else:
    print('----save file in notebook mode----')
    template.save("notebook.docx")