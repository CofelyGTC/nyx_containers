"""
BIAC IMPORT FEEDBACK COMMENTS
====================================

This process listens to two queues, decodes the file included in each message in order to create feedback result or comments.

Listens to:
-------------------------------------

* /queue/BAC_FEEDBACK_DOCX
* /queue/BAC_FEEDBACK_XLSX

Sends:
-------------------------------------

* /queue/BAC_FEEDBACK_RETURNMAIL_DOCX
* /queue/BAC_FEEDBACK_RETURNMAIL_XLSX


Collections:
-------------------------------------

* **biac_feedback_comment** (Raw Data)
* **biac_feedback_result** (Raw Data)
* **biac_feedback_status** (Raw Data)


VERSION HISTORY
===============

* 30 Jul 2019 0.0.1 **VME** First commit
* 31 Jul 2019 0.0.2 **VME** replacing key by the real key and adding the title. Adding user info
* 06 Aug 2019 0.0.3 **AMA** XLSX added for score
* 08 Aug 2019 0.0.4 **AMA** Feedback status added
* 22 Aug 2019 0.0.5 **AMA** Fix a logging bug
* 26 Aug 2019 0.0.6 **AMA** Fix a bug in the report entity / matching
* 24 Sep 2019 0.0.7 **AMA** Message localized in NL
* 23 Oct 2019 0.0.7 **AMA** Lot 4 Feedback result works
* 29 Oct 2019 1.0.0 **AMA** The system accepts KPI that are not numerical
* 04 Nov 2019 1.0.1 **AMA** Can read lot 5 excel files
* 05 Nov 2019 1.0.1 **AMA** Can read lot 5 docx files
* 21 Nov 2019 1.1.0 **AMA** Can read lot 6/7 excel files
* 25 Nov 2019 1.2.0 **AMA** Can read lot 6/7 docx files
* 27 Nov 2019 1.3.0 **AMA** Removed one line of lot 6 and 7
* 03 Dec 2019 1.3.2 **VME** Fix bug with the user 
* 05 Dec 2019 1.3.3 **VME** Fix bug multiple comments in same paragraph
* 09 Dec 2019 1.4.0 **AMA** Fix the format of the KPI
* 12 Dec 2019 1.4.1 **AMA** Force lot 1 all et lot 3 all to the appropriate contract.
"""  
import re
import json
import time
import uuid
import base64
import locale
import platform
import threading
import os,logging
import numpy as np

import pandas as pd

from docx import Document
from functools import wraps
from datetime import datetime
from datetime import timedelta
from elastic_helper import es_helper 
from amqstompclient import amqstompclient
from logging.handlers import TimedRotatingFileHandler
from logstash_async.handler import AsynchronousLogstashHandler
from elasticsearch import Elasticsearch as ES, RequestsHttpConnection as RC



MODULE  = "BIAC_FEEDBACK_COMMENTS_IMPORTER"
VERSION = "1.4.2"
QUEUE   = ["BAC_FEEDBACK_XLSX","BAC_FEEDBACK_DOCX"]




def log_message(message):
    global conn

    message_to_send={
        "message":message,
        "@timestamp":datetime.now().timestamp() * 1000,
        "module":MODULE,
        "version":VERSION
    }
    logger.info("LOG_MESSAGE")
    logger.info(message_to_send)
    conn.send_message("/queue/NYX_LOG",json.dumps(message_to_send))

def set_xlsx_status(es,user,reporttype,reportdate):
    key=(reportdate.strftime("%d%m%Y")+"_"+reporttype).lower().replace(" ","").replace(")","").replace("(","_")
    logger.info("KEY="+key)
    try:
        statusobj=es.get(index="biac_feedback_status",id=key,doc_type="doc")["_source"]
        statusobj["xlsx"]=True
        statusobj["user"]=user
        statusobj["xlsx_date"]=datetime.utcnow().isoformat()+"Z"
    except:
        statusobj={"user":user,"reporttype":reporttype,"reportdate":reportdate,"xlsx":True,"docx":False,"sent":False,"creation_date":datetime.now()}
    res=es.index(index="biac_feedback_status",doc_type="doc",body=statusobj,id=key)
    logger.info(res)

def set_docx_status(es,user,reporttype,reportdate):
    key=(reportdate.strftime("%d%m%Y")+"_"+reporttype).lower().replace(" ","").replace(")","").replace("(","_")
    logger.info("KEY="+key)
    try:
        statusobj=es.get(index="biac_feedback_status",id=key,doc_type="doc")["_source"]
        statusobj["docx"]=True
        statusobj["user"]=user
        statusobj["docx_date"]=datetime.utcnow().isoformat()+"Z"
    except:
        statusobj={"user":user,"reporttype":reporttype,"reportdate":reportdate,"xlsx":False,"docx":True,"sent":False,"creation_date":datetime.now()}
    es.index(index="biac_feedback_status",doc_type="doc",body=statusobj,id=key)

    


def getEntityObj(es):

    res=es.search(index="biac_entity",body={"size":100}) 
    entityObj = {}
    for i in res['hits']['hits']:
        if 'title' in i['_source']:
            #print(i['_source']['title'])
            #print('  ->'+i['_source']['contract'])

            entityObj[i['_source']['title']] = {
                'contract': i['_source']['contract']
            }

            if 'lot' in i['_source']:
                #print('  ->'+str(i['_source']['lot']))
                entityObj[i['_source']['title']]['lot']=str(i['_source']['lot'])
            if 'technic' in i['_source']:
                #print('  ->'+str(i['_source']['technic']))
                entityObj[i['_source']['title']]['technic']=str(i['_source']['technic'])
            else:
                entityObj[i['_source']['title']]['technic']=''
            
            if 'key' in i['_source']:
                #print('  ->'+str(i['_source']['technic']))
                entityObj[i['_source']['title']]['key']=str(i['_source']['key'])
            else:
                entityObj[i['_source']['title']]['key']=''


    return entityObj

def getEntityObjXLS(es,entity):

    res=es.search(index="biac_entity",body={"size":100}) 
    for row in res["hits"]["hits"]:
        if row["_source"]["key"]==entity:
            return row["_source"]
    return None

def cleankpi(x):
    try:
        return int(x)
    except:
        return x

################################################################################
def messageReceived(destination,message,headers):
    logger.info(">>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>")
    logger.info("Message Received:"+destination)

    if "XLSX" in destination:
        messageReceivedXLSX(destination,message,headers)      
    elif "DOCX" in destination:
        messageReceivedDOCX(destination,message,headers)
    else:
        logger.error("Unknown destination %s" %(destination))



##################################################################################
def messageReceivedXLSX(destination,message,headers):
    global es
    starttime = time.time()
    logger.info("==> "*10)
    logger.info("XLS Message Received %s" % destination)
    logger.info(headers)

    if "CamelSplitAttachmentId" in headers:
        headers["file"] = headers["CamelSplitAttachmentId"]

    if "file" in headers:
        logger.info("File:%s" %headers["file"])
        log_message("Import of file [%s] started." % headers["file"])



    try:
        xlsbytes = base64.b64decode(message)
        f = open('./tmp/excel.xlsx', 'wb')
        f.write(xlsbytes)
        f.close()

        if 'user' in headers:
            user_obj = json.loads(headers['user'])
            user = user_obj['firstname'] + ' ' + user_obj['lastname']
            user_id = user_obj['id']
        
        if 'From' in headers:
            mailfrom=headers['From'].split("<")[1][:-1]
            user_obj=es.get(index="nyx_user",id=mailfrom,doc_type="doc")
            user_id = user_obj['_id']
            user_obj= user_obj["_source"]
            user = user_obj['firstname'] + ' ' + user_obj['lastname']
            

        logger.info(user_obj)
        logger.info(user)
        logger.info(user_id)

        dftop=pd.read_excel("./tmp/excel.xlsx",skiprows=0)
        if dftop.columns[7]=="LOT5":
            df=pd.read_excel("./tmp/excel.xlsx",skiprows=8)
            dforg=df.copy()
            df=df[[df.columns[0],df.columns[7]]]
            df.columns=["KPI","NEGO"]
            nscore=df[pd.notnull(df['KPI']) & pd.notnull(df['NEGO'])]
            print(nscore)
            reporttype=dftop.columns[7]
            if(reporttype=="LOT5"):
                reporttype="Lot5"

            reportdate=datetime.strptime(dftop.columns[6],"%d/%m/%Y")
            entity={"lot":5,"contract":"Lot5","title":"Lot5"}
        elif dftop.columns[7]=="LOT6":
            df=pd.read_excel("./tmp/excel.xlsx",skiprows=7)
            dforg=df.copy()
            df=df[[df.columns[0],df.columns[11]]]
            df.columns=["KPI","NEGO"]
            nscore=df[pd.notnull(df['KPI']) & pd.notnull(df['NEGO'])]
            print(nscore)
            reporttype=dftop.columns[7]
            if(reporttype=="LOT6"):
                reporttype="Lot6"

            reportdate=datetime.strptime(dftop.columns[6],"%d/%m/%Y")
            entity={"lot":6,"contract":"Lot6","title":"Lot6"}
        elif dftop.columns[7]=="LOT7":
            df=pd.read_excel("./tmp/excel.xlsx",skiprows=7)
            dforg=df.copy()
            df=df[[df.columns[0],df.columns[11]]]
            df.columns=["KPI","NEGO"]
            nscore=df[pd.notnull(df['KPI']) & pd.notnull(df['NEGO'])]
            print(nscore)
            reporttype=dftop.columns[7]
            if(reporttype=="LOT7"):
                reporttype="Lot7"

            reportdate=datetime.strptime(dftop.columns[6],"%d/%m/%Y")
            entity={"lot":7,"contract":"Lot7","title":"Lot7"}
        else:
            df=pd.read_excel("./tmp/excel.xlsx",skiprows=7)
            dforg=df.copy()
            df=df[[df.columns[1],df.columns[10]]]
            df.columns=["KPI","NEGO"]

            nscore=df[pd.notnull(df['KPI']) & pd.notnull(df['NEGO'])]
            print(nscore)
            
            df=pd.read_excel("./tmp/excel.xlsx")

    #        print(df)
            try:
                reporttype=df.columns[8]
                reportdate=datetime.strptime(df.columns[7],"%d/%m/%Y")
            except: # LOT 4
                logger.info("Unable to decode report date")
                print(df.columns)
                reporttype=df.columns[7]
                reportdate=datetime.strptime(df.columns[6],"%d/%m/%Y")

                dforg=dforg[[dforg.columns[2],dforg.columns[14]]]
                dforg.columns=["KPI","NEGO"]
                #dforg=dforg[dforg['NEGO']!="/"]

                nscore=dforg[(pd.notnull(dforg['KPI'])) & (pd.notnull(dforg['NEGO']))]
                print(nscore)
            
            reporttype=reporttype.replace("Lot4 (DNB)","Lot4 (BACDNB)")
            entity=getEntityObjXLS(es,reporttype)


            

        maanden=['Januari',
            'Februari',
            'Maart',
            'April',
            'Mei',
            'Juni',
            'Juli',
            'Augustus',
            'September',
            'Oktober',
            'November',
            'December']
        

        reportdateNL=maanden[reportdate.month-1]


        logger.info(entity)
        
        results=[{"kpi":x[1][0],"result":x[1][1]} for x in nscore.iterrows()]

        
        scorebody="".join(["<li><b>KPI:</b>"+str(x["kpi"])+" <b>Resultaat:</b>"+str(x["result"])+ "</li>" for x in results])
        scorebody="<ul>"+scorebody+"</ul>"

        returnmail={"mail":user_id, "results":results,"user":user,'reportdate_nl': reportdateNL,"reportdate":reportdate.strftime("%d/%m/%Y"),"scorebody":scorebody,"entity":entity}
        logger.info(json.dumps(returnmail))

        dict_comment=[]

        

        for result in results:
            obj = {
                'key': reporttype,
                'title': entity["title"],
                'lot': entity["lot"],
                'contract': entity["contract"],
                'technic': entity.get("technic",""),
                'report_date': reportdate,                
                'creation_date': datetime.now(),
                'kpi': cleankpi(result["kpi"]),
                'result': result["result"],
                'user': user,
                'user_id': user_id,
            }

            dict_comment.append(obj)

        df_score=pd.DataFrame(dict_comment)     

        logger.info(df_score)

        if len(df_score) > 0:

            df_score['creation_date']=df_score['creation_date'].dt.tz_localize(tz='Europe/Paris')                
            df_score['report_date']=df_score['report_date'].dt.tz_localize(tz='Europe/Paris')                
            df_score['_index']='biac_feedback_result'
            #df_score['_id']=(reportdate.strftime("%d%m%Y")+"_"+reporttype+"_"+df_score["kpi"])#.lower().replace(" ","").replace(")","").replace("(","_")
            df_score['_id']=df_score["kpi"].apply(lambda x:(reportdate.strftime("%d%m%Y")+"_"+reporttype+"_"+str(x)).lower().replace(" ","").replace(")","").replace("(","_"))

            logger.info(df_score)
            es_helper.dataframe_to_elastic(es, df_score)  

        set_xlsx_status(es,user,reporttype,reportdate)
        conn.send_message("/queue/BAC_FEEDBACK_RETURNMAIL_XLSX",json.dumps(returnmail))

    except Exception as e:
        endtime = time.time()
        logger.error(e,exc_info=e)
        log_message("Import of file [%s] failed. Duration: %d Exception: %s." % (headers["file"],(endtime-starttime),str(e)))        


    endtime = time.time()    
    try:
        log_message("Import of file [%s] finished. Duration: %d Records: %d." % (headers["file"],(endtime-starttime),df_comment.shape[0]))         
    except:
        log_message("Import of file [%s] finished. Duration: %d." % (headers["file"],(endtime-starttime)))    




#================================================================
def messageReceivedDOCX(destination,message,headers):
    global es
    starttime = time.time()
    logger.info("==> "*10)
    logger.info("DOC Message Received %s" % destination)
    logger.info(headers)

    df_comment=pd.DataFrame()

    entityObj = getEntityObj(es)
    logger.info(entityObj)

    if "CamelSplitAttachmentId" in headers:
        headers["file"] = headers["CamelSplitAttachmentId"]

    if "file" in headers:
        logger.info("File:%s" %headers["file"])
        log_message("Import of file [%s] started." % headers["file"])



    docbytes = base64.b64decode(message)
    f = open('./tmp/word.docx', 'wb')
    f.write(docbytes)
    f.close()

    try:       
        doc = Document('./tmp/word.docx')

        dict_comment = []

        user = ''
        user_id = ''

        if 'user' in headers:
            user_obj = json.loads(headers['user'])
            user = user_obj['firstname'] + ' ' + user_obj['lastname']
            user_id = user_obj['id']

        if 'From' in headers:
            mailfrom=headers['From'].split("<")[1][:-1]
            user_obj=es.get(index="nyx_user",id=mailfrom,doc_type="doc")
            user_id = user_obj['_id']
            user_obj= user_obj["_source"]
            user = user_obj['firstname'] + ' ' + user_obj['lastname']
    
        logger.info(user)
        

        lot=0
        contract=''
        technic=''
        key=''
        report_date = None
        title=''

        for paragraph in doc.paragraphs:
            if paragraph.text.strip() != '' and paragraph.text.strip() != '\\n' :
                print(">>>>"+paragraph.text+"<<<<<")
                finalp=paragraph.text.strip()
                if finalp in entityObj:
                    lot=entityObj[finalp]['lot']
                    contract=entityObj[finalp]['contract']
                    technic=entityObj[finalp]['technic']
                    key=entityObj[finalp]['key'] 
                    title=finalp                
                    
                
                regex = 'KPI [a-zA-Z]{3,10} [0-9]{4}'
                x = re.search(regex, paragraph.text.strip())
                
                if x is not None:
                    report_date = datetime.strptime(paragraph.text, 'KPI %B %Y')

        logger.info('key      : '+key)
        logger.info('title    : '+title)
        logger.info('lot      : '+str(lot))
        logger.info('contract : '+contract)
        logger.info('technic  : '+technic)
        logger.info('date     : '+str(report_date))
        logger.info('user     : '+str(user))
        logger.info('user_id  : '+str(user_id))

        maanden=['Januari',
            'Februari',
            'Maart',
            'April',
            'Mei',
            'Juni',
            'Juli',
            'Augustus',
            'September',
            'Oktober',
            'November',
            'December']
        

        reportdateNL=maanden[report_date.month-1]

        results=[]

        if key != '':
            for paragraph in doc.paragraphs:
                if paragraph.text.strip() != '' and paragraph.text.strip() != '\\n' :
                    regex = '@([kK][pP][iI])? *([0-9]{1,3}[a-zA-Z]?) *:(.*)'
                    print(">>>>"+paragraph.text.strip())
                    matches = re.findall(regex, paragraph.text.strip())
                    if len(matches) > 0:
                        
                        for x in matches:
                            kpi = x[1]
                            comment = x[2].strip()
                            logger.info('     KPI COMMENT: '+kpi)
                            logger.info('         COMMENT: '+comment)

                            obj = {
                                'key': key.replace("Lot3 (All)","Lot3 (BACEXT)").replace("Lot1 (All)","Lot1 (BACHEA)"),
                                'title': title,
                                'lot': lot,
                                'contract': contract,
                                'technic': technic,
                                'report_date': report_date,
                                'creation_date': datetime.now(),
                                'kpi': kpi,
                                'comment': comment,
                                'user': user,
                                'user_id': user_id,
                            }
                            results.append({'kpi': kpi,
                            'comment': comment})
                            dict_comment.append(obj)

        df_comment=pd.DataFrame(dict_comment)     

        logger.info(df_comment)

        if len(df_comment) > 0:

            df_comment['creation_date']=df_comment['creation_date'].dt.tz_localize(tz='Europe/Paris')                
            df_comment['report_date']=df_comment['report_date'].dt.tz_localize(tz='Europe/Paris')                
            df_comment['_index']='biac_feedback_comment'

            es_helper.dataframe_to_elastic(es, df_comment)

            scorebody="".join(["<li><b>KPI:</b>"+str(x["kpi"])+" <b>Commentaar:</b>"+str(x["comment"])+ "</li>" for x in results])
            scorebody="<ul>"+scorebody+"</ul>"
            returnmail={"mail":user_id, "results":results,"user":user,'reportdate_nl': reportdateNL,"reportdate":report_date.strftime("%d/%m/%Y"),"scorebody":scorebody,"title":title}
            logger.info(json.dumps(returnmail))

            set_docx_status(es,user,key,report_date)

            conn.send_message("/queue/BAC_FEEDBACK_RETURNMAIL_DOCX",json.dumps(returnmail))
        else:
            returnmail={"mail":user_id, "results":results,"user":user,"reportdate":report_date.strftime("%d/%m/%Y"),'reportdate_nl': reportdateNL,"scorebody":"<div>- No Remarks Found</div>","title":title}
            set_docx_status(es,user,key,report_date)

            conn.send_message("/queue/BAC_FEEDBACK_RETURNMAIL_DOCX",json.dumps(returnmail))

    except Exception as e:
        endtime = time.time()
        logger.error(e,exc_info=e)
        log_message("Import of file [%s] failed. Duration: %d Exception: %s." % (headers["file"],(endtime-starttime),str(e)))        


    endtime = time.time()    
    try:
        log_message("Import of file [%s] finished. Duration: %d Records: %d." % (headers["file"],(endtime-starttime),df_comment.shape[0]))         
    except:
        log_message("Import of file [%s] finished. Duration: %d." % (headers["file"],(endtime-starttime)))    
    
        

if __name__ == '__main__':    
    locale.setlocale(locale.LC_TIME, "nl_BE")

    logging.basicConfig(level=logging.INFO,format='%(asctime)s %(levelname)s %(module)s - %(funcName)s: %(message)s', datefmt="%Y-%m-%d %H:%M:%S")
    logger = logging.getLogger()

    lshandler=None

    if os.environ["USE_LOGSTASH"]=="true":
        logger.info ("Adding logstash appender")
        lshandler=AsynchronousLogstashHandler("logstash", 5001, database_path='logstash_test.db')
        lshandler.setLevel(logging.ERROR)
        logger.addHandler(lshandler)

    handler = TimedRotatingFileHandler("logs/"+MODULE+".log",
                                    when="d",
                                    interval=1,
                                    backupCount=30)

    logFormatter = logging.Formatter('%(asctime)s.%(msecs)03d %(levelname)s %(module)s - %(funcName)s: %(message)s')
    handler.setFormatter( logFormatter )
    logger.addHandler(handler)

    logger.info("==============================")
    logger.info("Starting: %s" % MODULE)
    logger.info("Module:   %s" %(VERSION))
    logger.info("==============================")


    #>> AMQC
    server={"ip":os.environ["AMQC_URL"],"port":os.environ["AMQC_PORT"]
                    ,"login":os.environ["AMQC_LOGIN"],"password":os.environ["AMQC_PASSWORD"]}
    logger.info(server)                
    conn=amqstompclient.AMQClient(server
        , {"name":MODULE,"version":VERSION,"lifesign":"/topic/NYX_MODULE_INFO"},QUEUE,callback=messageReceived)
    #conn,listener= amqHelper.init_amq_connection(activemq_address, activemq_port, activemq_user,activemq_password, "RestAPI",VERSION,messageReceived)
    connectionparameters={"conn":conn}

    #>> ELK
    es=None
    logger.info (os.environ["ELK_SSL"])

    if os.environ["ELK_SSL"]=="true":
        host_params = {'host':os.environ["ELK_URL"], 'port':int(os.environ["ELK_PORT"]), 'use_ssl':True}
        es = ES([host_params], connection_class=RC, http_auth=(os.environ["ELK_LOGIN"], os.environ["ELK_PASSWORD"]),  use_ssl=True ,verify_certs=False)
    else:
        host_params="http://"+os.environ["ELK_URL"]+":"+os.environ["ELK_PORT"]
        es = ES(hosts=[host_params])


    logger.info("AMQC_URL          :"+os.environ["AMQC_URL"])
    while True:
        time.sleep(5)
        try:            
           
            variables={"platform":"_/_".join(platform.uname()),"icon":"comments"}
            conn.send_life_sign(variables=variables)
        except Exception as e:
            logger.error("Unable to send life sign.")
            logger.error(e)
    #app.run(threaded=True,host= '0.0.0.0')
